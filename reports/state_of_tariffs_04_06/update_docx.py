"""
Update state_of_tariffs.docx with new 04-06 numbers.

Copies the 04-02 docx as a starting point, then applies targeted text
replacements for all quantitative results.

Usage: python reports/state_of_tariffs_04_06/update_docx.py
"""

import csv
import os
import shutil
from docx import Document

DIR = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(os.path.dirname(DIR))
SRC_DOCX = os.path.join(REPO, 'reports', 'state_of_tariffs_04_02', 'state_of_tariffs.docx')
OUT_DOCX = os.path.join(DIR, 'state_of_tariffs.docx')

TEMP_DIR = os.path.join(REPO, 'output', '2026-04-06', 'results')
PERM_DIR = os.path.join(REPO, 'output', '2026-04-06_s122perm', 'results')


def read_key_results(path):
    with open(path, newline='') as f:
        rows = list(csv.DictReader(f))
    return {r['metric']: float(r['value']) for r in rows}

def read_csv_dict(path):
    with open(path, newline='') as f:
        return list(csv.DictReader(f))


temp = read_key_results(os.path.join(TEMP_DIR, 'key_results.csv'))
perm = read_key_results(os.path.join(PERM_DIR, 'key_results.csv'))

temp_sectors = {r['sector']: float(r['output_change_pct'])
                for r in read_csv_dict(os.path.join(TEMP_DIR, 'sector_effects.csv'))}
perm_sectors = {r['sector']: float(r['output_change_pct'])
                for r in read_csv_dict(os.path.join(PERM_DIR, 'sector_effects.csv'))}

temp_dist = sorted(read_csv_dict(os.path.join(TEMP_DIR, 'distribution.csv')),
                   key=lambda x: int(x['decile']))
perm_dist = sorted(read_csv_dict(os.path.join(PERM_DIR, 'distribution.csv')),
                   key=lambda x: int(x['decile']))

GDP_2025 = 28000

# New 04-06 values (rounded for display)
NEW = {
    'presub_allin_perm': f'{perm["pre_sub_all_in_etr"]:.1f}',    # 12.2
    'postsub_allin_perm': f'{perm["pe_postsub_all_in_etr"]:.1f}', # 10.5
    'presub_allin_temp': f'{temp["pre_sub_all_in_etr"]:.1f}',     # 9.7
    'postsub_allin_temp': f'{temp["pe_postsub_all_in_etr"]:.1f}', # 8.2
    'presub_incr_perm': f'{perm["pre_sub_etr_increase"]:.1f}',    # 10.2
    'postsub_incr_perm': f'{perm["pe_postsub_etr_increase"]:.1f}',# 8.5
    'presub_price_temp': f'{temp["pre_sub_price_increase"]:.1f}',  # 0.7
    'postsub_price_temp': f'{temp["pe_postsub_price_increase"]:.1f}', # 0.5
    'presub_price_perm': f'{perm["pre_sub_price_increase"]:.1f}',  # 1.1
    'postsub_price_perm': f'{perm["pe_postsub_price_increase"]:.1f}', # 0.9
    'presub_hh_temp': f'{temp["pre_sub_per_hh_cost"]:,.0f}',      # 939
    'postsub_hh_temp': f'{temp["post_sub_per_hh_cost"]:,.0f}',    # 761
    'presub_hh_perm': f'{perm["pre_sub_per_hh_cost"]:,.0f}',      # 1,478
    'postsub_hh_perm': f'{perm["post_sub_per_hh_cost"]:,.0f}',    # 1,208
    'lr_gdp_temp': f'{abs(temp_sectors["overall_gdp"]):.2f}',      # 0.11
    'lr_gdp_perm': f'{abs(perm_sectors["overall_gdp"]):.2f}',      # 0.18
    'lr_gdp_dollars': str(round(abs(temp_sectors['overall_gdp']) / 100 * GDP_2025)),
    'mfg_temp': f'{temp_sectors["manufacturing"]:.1f}',            # 1.1
    'const_temp': f'{abs(temp_sectors["construction"]):.1f}',      # 2.5
    'mining_temp': f'{abs(temp_sectors["mining"]):.1f}',           # 1.0
    'conv_rev_temp_t': f'{temp["conventional_revenue_10yr"]/1000:.1f}',
    'dyn_rev_temp_t': f'{temp["dynamic_revenue_10yr"]/1000:.1f}',
    'conv_rev_perm_t': f'{perm["conventional_revenue_10yr"]/1000:.1f}',
    'dyn_rev_perm_t': f'{perm["dynamic_revenue_10yr"]/1000:.1f}',
    'dyn_effect_temp': f'{abs(temp["dynamic_effect_10yr"]):.0f}',
    'dist_d1_temp': f'{abs(float(temp_dist[0]["pct_of_income"])):.1f}',
    'dist_d10_temp': f'{abs(float(temp_dist[9]["pct_of_income"])):.1f}',
    'dist_d1_perm': f'{abs(float(perm_dist[0]["pct_of_income"])):.1f}',
    'dist_d10_perm': f'{abs(float(perm_dist[9]["pct_of_income"])):.1f}',
    'dist_d1_cost_temp': f'{abs(float(temp_dist[0]["cost_per_hh"])):,.0f}',
    'dist_d10_cost_temp': f'{abs(float(temp_dist[9]["cost_per_hh"])):,.0f}',
    'dist_d1_cost_perm': f'{abs(float(perm_dist[0]["cost_per_hh"])):,.0f}',
    'dist_d10_cost_perm': f'{abs(float(perm_dist[9]["cost_per_hh"])):,.0f}',
}

d1_v_d10 = abs(float(temp_dist[0]['pct_of_income'])) / abs(float(temp_dist[9]['pct_of_income']))
NEW['d1_ratio'] = f'{d1_v_d10:.0f}'


# ---- Helper: robust paragraph text replacement ----------------------------

def replace_in_para(para, old_text, new_text):
    """Replace old_text with new_text in paragraph. Tries single-run first,
    then falls back to paragraph-level reconstruction (loses mixed formatting)."""
    # Single-run replacement (preserves formatting)
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text, 1)
            return True

    # Cross-run fallback
    full = para.text
    if old_text not in full:
        return False

    new_full = full.replace(old_text, new_text, 1)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ''
    return True


# ---- Copy and open --------------------------------------------------------

shutil.copy2(SRC_DOCX, OUT_DOCX)
doc = Document(OUT_DOCX)
changes = []


# ---- Apply replacements per paragraph ------------------------------------

# OLD values from the 04-02 docx (after its update_docx.py ran)
# These are what we're searching for and replacing.

for i, para in enumerate(doc.paragraphs):
    full = para.text

    # == Date references ==
    if 'April 2, 2026' in full:
        replace_in_para(para, 'April 2, 2026', 'April 6, 2026')
        changes.append(f'P{i}: Date April 2 -> April 6')

    if 'April 2' in para.text and 'as of' in full.lower():
        replace_in_para(para, 'April 2', 'April 6')

    # == P3-equiv: ETR Summary ==
    # 04-02 docx has: "stands at 11.0%...fall to 8.2%...9.6% and 7.1%"
    if 'average effective tariff rate stands at' in full:
        replace_in_para(para, '11.0%', f'{NEW["presub_allin_perm"]}%')
        replace_in_para(para, '8.2%', f'{NEW["presub_allin_temp"]}%')
        replace_in_para(para, '9.6%', f'{NEW["postsub_allin_perm"]}%')
        replace_in_para(para, '7.1%', f'{NEW["postsub_allin_temp"]}%')
        # Update historical comparison
        replace_in_para(para, 'highest since 1943', 'highest since 1940')
        changes.append(f'P{i}: ETR summary updated')

    # == P4-equiv: Price Summary ==
    # 04-02: "between 0.5% and 0.6%...$650...$780..between 0.8% and 1.0%...$1,130...$1,340"
    elif 'ultimate price level impact will be between' in full:
        replace_in_para(para, 'between 0.5% and 0.6%',
                        f'between {NEW["postsub_price_temp"]}% and {NEW["presub_price_temp"]}%')
        replace_in_para(para, '$650', f'${NEW["postsub_hh_temp"]}')
        replace_in_para(para, '$780', f'${NEW["presub_hh_temp"]}')
        replace_in_para(para, 'between 0.8% and 1.0%',
                        f'between {NEW["postsub_price_perm"]}% and {NEW["presub_price_perm"]}%')
        replace_in_para(para, '$1,130', f'${NEW["postsub_hh_perm"]}')
        replace_in_para(para, '$1,340', f'${NEW["presub_hh_perm"]}')
        changes.append(f'P{i}: Price summary updated')

    # == P5-equiv: Macro Summary ==
    elif 'persistently' in full and 'billion annually' in full:
        replace_in_para(para, f'0.10%', f'{NEW["lr_gdp_temp"]}%')
        replace_in_para(para, '$27 billion', f'${NEW["lr_gdp_dollars"]} billion')
        replace_in_para(para, '0.16%', f'{NEW["lr_gdp_perm"]}%')
        changes.append(f'P{i}: Macro summary updated')

    # == P6-equiv: Sector Summary ==
    elif 'manufacturing output expands by' in full:
        replace_in_para(para, 'expands by 0.7%', f'expands by {NEW["mfg_temp"]}%')
        replace_in_para(para, 'contracts by 2.0%', f'contracts by {NEW["const_temp"]}%')
        replace_in_para(para, 'declines by 0.8%', f'declines by {NEW["mining_temp"]}%')
        changes.append(f'P{i}: Sector summary updated')

    # == P7-equiv: Fiscal Summary ==
    elif 'raise about $1.1 trillion over 2026' in full or \
         'raise about $1.3 trillion over 2026' in full:
        replace_in_para(para, '$1.1 trillion', f'${NEW["conv_rev_temp_t"]} trillion')
        replace_in_para(para, '$1.0 trillion', f'${NEW["dyn_rev_temp_t"]} trillion')
        replace_in_para(para, '$1.7 trillion', f'${NEW["conv_rev_perm_t"]} trillion')
        replace_in_para(para, '$1.6 trillion', f'${NEW["dyn_rev_perm_t"]} trillion')
        changes.append(f'P{i}: Fiscal summary updated')

    # == ETR Detail: Pre-sub ==
    elif 'percentage point increase in the US average effective tariff rate' in full and \
         'pre-substitution' in full.lower():
        replace_in_para(para, '9.0 percentage point', f'{NEW["presub_incr_perm"]} percentage point')
        replace_in_para(para, '11.0%', f'{NEW["presub_allin_perm"]}%')
        replace_in_para(para, 'highest since 1943', 'highest since 1940')
        changes.append(f'P{i}: ETR detail pre-sub updated')

    # == ETR Detail: Post-sub ==
    elif 'percentage point increase in the US average effective tariff rate' in full and \
         'post-substitution' in full.lower():
        replace_in_para(para, '7.6 percentage point', f'{NEW["postsub_incr_perm"]} percentage point')
        replace_in_para(para, '9.6%', f'{NEW["postsub_allin_perm"]}%')
        replace_in_para(para, '8.2%', f'{NEW["presub_allin_temp"]}%')
        replace_in_para(para, '7.1%', f'{NEW["postsub_allin_temp"]}%')
        changes.append(f'P{i}: ETR detail post-sub updated')

    # == Price Detail: Main para ==
    elif 'increase in consumer prices of' in full and 'short run' in full:
        replace_in_para(para, 'of 1.0%', f'of {NEW["presub_price_perm"]}%')
        replace_in_para(para, 'about 0.6%', f'about {NEW["presub_price_temp"]}%')
        replace_in_para(para, '$1,338', f'${NEW["presub_hh_perm"]}')
        replace_in_para(para, '$780', f'${NEW["presub_hh_temp"]}')
        changes.append(f'P{i}: Price detail main updated')

    # == Price Detail: Post-sub para ==
    elif 'post-substitution price increase settles at' in full:
        replace_in_para(para, 'at 0.5%', f'at {NEW["postsub_price_temp"]}%')
        replace_in_para(para, '$648', f'${NEW["postsub_hh_temp"]}')
        replace_in_para(para, 'roughly 0.8%', f'roughly {NEW["postsub_price_perm"]}%')
        replace_in_para(para, '$1,130', f'${NEW["postsub_hh_perm"]}')
        changes.append(f'P{i}: Price detail post-sub updated')

    # == Macro Detail ==
    elif '0.10% to 0.16% smaller' in full or \
         f'{NEW["lr_gdp_temp"]}% to' in full:
        replace_in_para(para, '0.10%', f'{NEW["lr_gdp_temp"]}%')
        replace_in_para(para, '0.16%', f'{NEW["lr_gdp_perm"]}%')
        changes.append(f'P{i}: Macro detail updated')

    # == Fiscal Detail: Expires ==
    elif 'total about $' in full and 'billion over the decade' in full:
        replace_in_para(para, '$90 billion', f'${NEW["dyn_effect_temp"]} billion')
        changes.append(f'P{i}: Fiscal detail dynamic effect updated')

    # == Fiscal Detail: Extended ==
    elif 'dynamic score would be roughly' in full:
        replace_in_para(para, '$1.6 trillion', f'${NEW["dyn_rev_perm_t"]} trillion')
        changes.append(f'P{i}: Fiscal detail extended updated')

    # == Distribution Detail ==
    elif 'first decile is about' in full and 'times that of the top' in full:
        replace_in_para(para, '2 times', f'{NEW["d1_ratio"]} times')
        # D1/D10 pct of income - temp
        replace_in_para(para, '0.8% versus 0.4%',
                        f'{NEW["dist_d1_temp"]}% versus {NEW["dist_d10_temp"]}%')
        # D1/D10 pct of income - perm
        replace_in_para(para, '1.6% versus 0.7%',
                        f'{NEW["dist_d1_perm"]}% versus {NEW["dist_d10_perm"]}%')
        # HH costs - temp
        replace_in_para(para, '$430', f'${NEW["dist_d1_cost_temp"]}')
        replace_in_para(para, '$1,810', f'${NEW["dist_d10_cost_temp"]}')
        # HH costs - perm
        replace_in_para(para, '$740', f'${NEW["dist_d1_cost_perm"]}')
        replace_in_para(para, '$3,100', f'${NEW["dist_d10_cost_perm"]}')
        changes.append(f'P{i}: Distribution updated')

    # == Revenue table paragraph (raise about) ==
    elif 'raise about' in full and 'conventionally scored' in full:
        replace_in_para(para, '$1.1 trillion', f'${NEW["conv_rev_temp_t"]} trillion')
        replace_in_para(para, '$90 billion', f'${NEW["dyn_effect_temp"]} billion')
        replace_in_para(para, '$1.0 trillion', f'${NEW["dyn_rev_temp_t"]} trillion')
        changes.append(f'P{i}: Revenue body text updated')

    # == Revenue extended paragraph ==
    elif 'conventional revenue would be' in full:
        replace_in_para(para, '$1.7 trillion', f'${NEW["conv_rev_perm_t"]} trillion')
        replace_in_para(para, '$1.6 trillion', f'${NEW["dyn_rev_perm_t"]} trillion')
        changes.append(f'P{i}: Revenue extended text updated')


# ---- Also update the "Changes Since" section heading ----
for i, para in enumerate(doc.paragraphs):
    if 'Changes Since the Last Report' in para.text:
        # The following paragraphs are the "Changes" content.
        # We note that the user should update this section manually.
        pass

    # Update the March 9 reference to April 2 in the changes section
    if 'March 9' in para.text and 'budgetlab.yale.edu' in para.text:
        replace_in_para(para, 'March 9', 'April 2')
        replace_in_para(para, 'march-9-2026', 'april-2-2026')
        changes.append(f'P{i}: Updated prior report reference to April 2')


# ---- Save ----------------------------------------------------------------

doc.save(OUT_DOCX)

print(f'Updated {OUT_DOCX}')
print(f'Source: {SRC_DOCX}')
print(f'\nChanges made ({len(changes)}):')
for c in changes:
    print(f'  {c}')

print(f'\n--- New values ---')
for k, v in NEW.items():
    print(f'  {k}: {v}')

print('\n*** MANUAL REVIEW NEEDED ***')
print('  - "Changes Since the Last Report" section content')
print('  - Section 232 policy table (metal overhaul, pharma)')
print('  - Historical year comparisons ("highest since X")')
print('  - Any cross-run text replacements that may have lost formatting')
